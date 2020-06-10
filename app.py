import time

from collections import OrderedDict

from docx import Document
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
import keras
import matplotlib.pyplot as plt
import base64
import numpy as np
import io

from PIL import Image
# from keras.models import Sequential
from pytesseract import pytesseract
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Dense, Dropout, LSTM
from tensorflow.keras.models import load_model
from tensorflow.keras.preprocessing.image import img_to_array
import cv2

app = Flask(__name__, template_folder='static')

CORS(app)

fonts_dict = {'Algerian': 'ALGER',
              'Arial': 'arial',
              'Calibri': 'calibri',
              'Cambria Math': 'cambria',
              'Candara': 'Candara',
              'Centaur': 'CENTAUR',
              'Gill Sans MT': 'GIL_____',
              'Tahoma': 'tahoma',
              'Times New Roman': 'times',
              'Verdana': 'verdana'}

classes = {0: 'Algerian',
           1: 'Arial',
           2: 'Calibri',
           3: 'Cambria Math',
           4: 'Candara',
           5: 'Centaur',
           6: 'Gill Sans MT',
           7: 'Tahoma',
           8: 'Times New Roman',
           9: 'Verdana'}


def get_model():
    global model
    model = load_model('font_recognization.h5')
    model.load_weights('weights.best.hdf5')
    print('model loaded')


def preprocess_image(image, target_size):
    image = np.resize(image, target_size)
    x = img_to_array(image)  # Numpy array with shape (300, 300, 3)
    x = x.reshape((1,) + x.shape)  # Numpy array with shape (1, 300,300, 3)

    # Rescale by 1/255
    x /= 255.0
    return x


# return paragraph in reverse order
def get_paragraph(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(gray, (7, 7), 0)
    thresh = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
    dilate = cv2.dilate(thresh, kernel, iterations=10)

    cnts = cv2.findContours(dilate, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]

    X_data = []

    for c in cnts:
        x, y, w, h = cv2.boundingRect(c)
        X_data.append(image[y:y + h, x:x + w])

    return X_data


def get_lines(img):
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # (2) threshold
    th, threshed = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)

    # (5) find and draw the upper and lower boundary of each lines
    hist = cv2.reduce(threshed, 1, cv2.REDUCE_AVG).reshape(-1)

    th = 0
    H, W = img.shape[:2]
    uppers = [y for y in range(H - 1) if hist[y] <= th and hist[y + 1] > th]
    lowers = [y for y in range(H - 1) if hist[y] > th and hist[y + 1] <= th]

    X_data = []

    for i in range(min(len(uppers), len(lowers))):
        X_data.append(img[uppers[i]:lowers[i], 0:W])

    return X_data


def get_word(img):
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    ## (2) threshold
    th, threshed = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)

    hist = cv2.reduce(threshed, 0, cv2.REDUCE_AVG).reshape(-1)

    th = 0
    H, W = img.shape[:2]
    left = [x for x in range(W - 1) if hist[x] <= th and hist[x + 1] > th]
    right = [x for x in range(W - 1) if hist[x] > th and hist[x + 1] <= th]

    X_data = []

    for i in range(min(len(left), len(right))):
        # cv2.imwrite(image_name[i], img[uppers[i]:lowers[i],0:W])
        X_data.append(img[0:H, left[i]:right[i]])

    return X_data


def create_document(font_text_dict):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font_text_dict = OrderedDict(font_text_dict)
    for text in font_text_dict:
        font.name = font_text_dict[text]
        paragraph = document.add_paragraph(text)
    document_name = time.strftime("%Y%m%d-%H%M%S") + ".doc"

    document.save(document_name)
    return document_name


def create_doc(text):
    document = Document()

    document.add_paragraph(text)

    document_name = time.strftime("%Y%m%d-%H%M%S") + ".doc"

    document.save(document_name)
    return document_name


# loading model
get_model()


@app.route('/')
def home():
    return render_template('predict.html')


@app.route('/get-document', methods=['POST'])
def predict():
    message = request.get_json(force=True)
    encoded = message['image']
    decoded = base64.b64decode(encoded)
    image = Image.open(io.BytesIO(decoded))
    if image.mode != 'RGB':
        image = image.convert('RGB')
    print(np.array(image).shape)
    paragraphs = get_paragraph(np.array(image))
    fonts_text_dict = {}

    for p in reversed(paragraphs):
        # get text in paragraph
        text_in_para = pytesseract.image_to_string(p)

        # get images of lines in paragraph
        lines_in_para = get_lines(np.array(p))
        print(len(lines_in_para))
        # get images of words in line
        if len(lines_in_para) == 0:
            processed_image = preprocess_image(p, target_size=(300, 300, 3))
        elif len(lines_in_para) == 1:
            processed_image = preprocess_image(lines_in_para[0], target_size=(300, 300, 3))
        else:
            words_in_line = get_word(lines_in_para[0])
            # ready image for model
            processed_image = preprocess_image(words_in_line[0], target_size=(300, 300, 3))

        # get font in paragraph
        fonts_in_para = classes[np.argmax(model.predict(processed_image), axis=1).tolist()[0]]
        fonts_text_dict.update({text_in_para: fonts_in_para})

    print('dictionary created')
    document_name = create_document(fonts_text_dict)
    # create_doc(pytesseract.image_to_string(image))

    response = {
        'document-name': document_name
    }

    return jsonify(response)


@app.route('/download/<document_name>')
def download_file(document_name):
    return send_file(document_name, mimetype='/text/doc/', as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
